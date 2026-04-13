from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# PDF layout constants. Units are in points (1/72 inch). These can be adjusted as needed
# for different formatting preferences.
PDF_MARGIN = 72
PDF_TEXT_SIZE = 12
PDF_TITLE_SIZE = 18
PDF_SECTION_SIZE = 14
PDF_LINE_HEIGHT = 18

PDF_FONT_STANDARD = "Helvetica"
PDF_FONT_STANDARD_BOLD = "Helvetica-Bold"
PDF_FONT_STANDARD_CJK = "STSong-Light"

_FONTS_DIR = Path(__file__).parent.parent / "assets" / "fonts"
_PDF_FONT_TTF_REGULAR = "TranscriberRegular"
_PDF_FONT_TTF_BOLD = "TranscriberBold"
_PDF_FONT_TTF_CJK = "TranscriberCJK"
_PDF_FONT_TTF_CJK_BOLD = "TranscriberCJKBold"



# This function takes the cleaned transcript, summary, and metadata, and writes them to JSON,
# DOCX, and PDF files in the specified output directory. It returns a dictionary with the paths
# to each of the generated files.
def write_outputs(
    *,
    output_dir: Path,
    stem: str,
    cleaned_transcript: str | None,
    summary: list[str],
    metadata: dict[str, Any],
    audio_filename: str | None = None,
    raw_transcript: str | None = None,
    title: str = "Audio Transcription Output",
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)

    json_path = output_dir / f"{stem}.json"
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"
    content = _build_export_content(
        audio_filename,
        cleaned_transcript,
        raw_transcript,
        summary,
        metadata,
        title,
    )
    json_path.write_text(json.dumps(content, indent=2, ensure_ascii=False), encoding="utf-8")
    _write_docx(docx_path, content)
    _write_pdf(pdf_path, content)

    return {
        "json": json_path,
        "docx": docx_path,
        "pdf": pdf_path,
    }


# This helper function constructs a structured content dictionary that includes the audio filename,
# cleaned transcript (or raw transcript if cleaned is unavailable),
# summary, and
# metadata.
def _build_export_content(
    audio_filename: str | None,
    cleaned_transcript: str | None,
    raw_transcript: str | None,
    summary: list[str],
    metadata: dict[str, Any],
    title: str = "Audio Transcription Output",
) -> dict[str, Any]:
    transcript = (cleaned_transcript or raw_transcript or "").strip() or "[Transcript unavailable]"
    generated_at = datetime.now().astimezone().strftime("%m-%d-%Y %H:%M:%S %Z")
    return {
        "title": title,
        "generated_at": generated_at,
        "audio_filename": audio_filename,
        "metadata": dict(metadata),
        "summary": list(summary),
        "transcript_lines": transcript.splitlines() or [transcript],
    }


# The following two functions, _write_docx and _write_pdf,
# are responsible for generating the DOCX and PDF files respectively.
# They take the structured content dictionary and format it appropriately for each file type, including headings,
# paragraphs, and
# metadata sections.
def _write_docx(path: Path, content: dict[str, Any]) -> None:
    document = Document()
    document.add_heading(content["title"], level=1)

    if content["audio_filename"]:
        document.add_paragraph(f"Source Audio: {content['audio_filename']}")
    document.add_paragraph(f"Generated: {content['generated_at']}")


    document.add_heading("Summary", level=2)
    summary: list[str] = content["summary"]
    if summary:
        for bullet in summary:
            document.add_paragraph(f"- {bullet}")
    else:
        document.add_paragraph("[No summary available]")

    document.add_heading("Transcript", level=2)
    transcript_lines: list[str] = content["transcript_lines"]
    for line in transcript_lines:
        document.add_paragraph(line)

    document.save(str(path))


# The _write_pdf function creates a PDF file with a structured layout that includes the title,
# metadata,
# summary, and
# transcript.

def _write_pdf(path: Path, content: dict[str, Any]) -> None:
    body_font, heading_font = _resolve_pdf_fonts(content)
    styles = _build_pdf_styles(body_font=body_font, heading_font=heading_font)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=PDF_MARGIN,
        rightMargin=PDF_MARGIN,
        topMargin=PDF_MARGIN,
        bottomMargin=PDF_MARGIN,
        pageCompression=0,
    )

    story: list[Any] = []
    story.append(Paragraph(_escape_pdf_text(content["title"]), styles["title"]))

    if content["audio_filename"]:
        story.append(Paragraph(_escape_pdf_text(f"Source Audio: {content['audio_filename']}"), styles["body"]))
    story.append(Paragraph(_escape_pdf_text(f"Generated: {content['generated_at']}"), styles["body"]))

    story.append(Spacer(1, PDF_LINE_HEIGHT / 2))
    story.append(Paragraph("Summary", styles["section"]))

    summary: list[str] = content["summary"]
    if summary:
        for bullet in summary:
            story.append(Paragraph(_escape_pdf_text(f"- {bullet}"), styles["body"]))
    else:
        story.append(Paragraph("[No summary available]", styles["body"]))

    story.append(Spacer(1, PDF_LINE_HEIGHT / 2))
    story.append(Paragraph("Transcript", styles["section"]))

    transcript_lines: list[str] = content["transcript_lines"]
    for line in transcript_lines:
        text = line if line.strip() else " "
        story.append(Paragraph(_escape_pdf_text(text), styles["body"]))

    doc.build(
        story,
        onFirstPage=lambda canvas, _doc: _set_pdf_metadata(canvas, content["title"]),
        onLaterPages=lambda canvas, _doc: _set_pdf_metadata(canvas, content["title"]),
    )


def _set_pdf_metadata(canvas: Any, title: str) -> None:
    canvas.setTitle(title)
    canvas.setAuthor("ai-audio-transcriber")
    canvas.setCreator("ai-audio-transcriber")
    canvas.setPageCompression(0)


def _build_pdf_styles(*, body_font: str, heading_font: str) -> dict[str, ParagraphStyle]:
    base_styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            name="TranscriberTitle",
            parent=base_styles["Heading1"],
            fontName=heading_font,
            fontSize=PDF_TITLE_SIZE,
            leading=PDF_LINE_HEIGHT + 3,
            spaceBefore=PDF_LINE_HEIGHT / 2,
            spaceAfter=PDF_LINE_HEIGHT / 2,
            wordWrap="CJK",
        ),
        "section": ParagraphStyle(
            name="TranscriberSection",
            parent=base_styles["Heading2"],
            fontName=heading_font,
            fontSize=PDF_SECTION_SIZE,
            leading=PDF_LINE_HEIGHT + 1,
            spaceBefore=PDF_LINE_HEIGHT / 2,
            spaceAfter=PDF_LINE_HEIGHT / 3,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            name="TranscriberBody",
            parent=base_styles["BodyText"],
            fontName=body_font,
            fontSize=PDF_TEXT_SIZE,
            leading=PDF_LINE_HEIGHT,
            spaceAfter=2,
            wordWrap="CJK",
        ),
    }


def _content_has_cjk(content: dict[str, Any]) -> bool:
    parts: list[str] = [content.get("title", "")]
    parts.extend(content.get("summary", []))
    parts.extend(content.get("transcript_lines", []))
    text = " ".join(parts)
    return any(0x4E00 <= ord(c) <= 0x9FFF or 0x3400 <= ord(c) <= 0x4DBF for c in text)


def _try_register_cjk_ttf_fonts() -> bool:
    regular_path = _FONTS_DIR / "font-cjk-regular.ttf"
    bold_path = _FONTS_DIR / "font-cjk-bold.ttf"
    if not regular_path.exists():
        return False
    try:
        if _PDF_FONT_TTF_CJK not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(_PDF_FONT_TTF_CJK, str(regular_path)))
        bold_name = _PDF_FONT_TTF_CJK_BOLD if bold_path.exists() else _PDF_FONT_TTF_CJK
        if bold_path.exists() and bold_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
        return True
    except Exception:
        return False


def _resolve_pdf_fonts(content: dict[str, Any]) -> tuple[str, str]:
    if _content_has_cjk(content):
        if _try_register_cjk_ttf_fonts():
            bold_name = _PDF_FONT_TTF_CJK_BOLD if (_FONTS_DIR / "font-cjk-bold.ttf").exists() else _PDF_FONT_TTF_CJK
            return _PDF_FONT_TTF_CJK, bold_name
        if _try_register_cid_font(PDF_FONT_STANDARD_CJK):
            return PDF_FONT_STANDARD_CJK, PDF_FONT_STANDARD_CJK
    if _try_register_ttf_fonts():
        return _PDF_FONT_TTF_REGULAR, _PDF_FONT_TTF_BOLD
    return PDF_FONT_STANDARD, PDF_FONT_STANDARD_BOLD


def _try_register_ttf_fonts() -> bool:
    regular_path = _FONTS_DIR / "font-regular.ttf"
    bold_path = _FONTS_DIR / "font-bold.ttf"
    if not regular_path.exists() or not bold_path.exists():
        return False
    try:
        if _PDF_FONT_TTF_REGULAR not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(_PDF_FONT_TTF_REGULAR, str(regular_path)))
        if _PDF_FONT_TTF_BOLD not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(_PDF_FONT_TTF_BOLD, str(bold_path)))
        return True
    except Exception:
        return False


def _try_register_cid_font(font_name: str) -> bool:
    if font_name in pdfmetrics.getRegisteredFontNames():
        return True
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
        return True
    except Exception:
        return False


def _escape_pdf_text(text: str) -> str:
    return escape(text)
